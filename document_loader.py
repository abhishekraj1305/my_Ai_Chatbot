"""Recursive document loading for PDF, DOCX, TXT, and Markdown files."""

from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

from pypdf import PdfReader

from utils.text_cleaner import clean_text

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - handled at runtime with a readable log
    DocxDocument = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class LoadedDocument:
    text: str
    metadata: dict


def _log(message: str) -> None:
    print(f"[document_loader] {message}")


def scan_documents(data_dir: str) -> List[Path]:
    """Return supported files under data_dir recursively."""
    root = Path(data_dir)
    if not root.exists():
        _log(f"Data folder not found: {root}")
        return []

    files = [
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    _log(f"Documents found: {len(files)}")
    return files


def _base_metadata(path: Path) -> dict:
    return {
        "file_name": path.name,
        "file_path": str(path),
        "document_type": path.suffix.lower().lstrip("."),
    }


def _load_pdf(path: Path) -> Iterable[LoadedDocument]:
    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if not text:
            continue

        metadata = _base_metadata(path)
        metadata["page_number"] = index
        yield LoadedDocument(text=text, metadata=metadata)


def _load_docx(path: Path) -> Optional[LoadedDocument]:
    if DocxDocument is None:
        _log("python-docx is not installed; skipping DOCX files.")
        return None

    document = DocxDocument(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = clean_text("\n".join(parts))
    if not text:
        return None

    return LoadedDocument(text=text, metadata=_base_metadata(path))


def _load_text_file(path: Path) -> Optional[LoadedDocument]:
    text = path.read_text(encoding="utf-8", errors="replace")
    text = clean_text(text)
    if not text:
        return None

    return LoadedDocument(text=text, metadata=_base_metadata(path))


def load_documents(data_dir: str) -> List[LoadedDocument]:
    """Load all supported documents with metadata and readable error logs."""
    loaded: List[LoadedDocument] = []
    files = scan_documents(data_dir)

    if not files:
        _log("No supported documents found. Add PDF, DOCX, TXT, or MD files.")
        return loaded

    for path in files:
        try:
            suffix = path.suffix.lower()
            before = len(loaded)

            if suffix == ".pdf":
                loaded.extend(_load_pdf(path))
            elif suffix == ".docx":
                doc = _load_docx(path)
                if doc:
                    loaded.append(doc)
            elif suffix in {".txt", ".md"}:
                doc = _load_text_file(path)
                if doc:
                    loaded.append(doc)

            added = len(loaded) - before
            if added:
                _log(f"Loaded {path.name} ({added} record{'s' if added != 1 else ''})")
            else:
                _log(f"Skipped empty file: {path.name}")
        except Exception as exc:
            _log(f"Could not load {path.name}: {exc}")

    _log(f"Documents loaded: {len(loaded)}")
    return loaded


if __name__ == "__main__":
    from config import DATA_DIR

    load_documents(DATA_DIR)

